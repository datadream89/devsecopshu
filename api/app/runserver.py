from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def iter_block_items(parent):
    from docx.oxml.ns import qn
    parent_elm = parent.element if hasattr(parent, "element") else parent

    for child in parent_elm.iterchildren():
        if child.tag == qn('w:p'):
            yield 'paragraph', child
        elif child.tag == qn('w:tbl'):
            yield 'table', child

def extract_docx_structure(doc_path):
    doc = Document(doc_path)

    paragraphs = doc.paragraphs
    tables = doc.tables
    para_idx = 0
    table_idx = 0

    output = []

    def is_bold(run): return run.bold is True
    def is_underline(run): return run.underline is True
    def is_bold_and_underline(run): return is_bold(run) and is_underline(run)

    for block_type, elm in iter_block_items(doc):
        print(f"Block type: {block_type}")  # DEBUG

        if block_type == 'paragraph':
            if para_idx >= len(paragraphs):
                print("Paragraph index out of range!")
                break
            para = paragraphs[para_idx]
            para_idx += 1
            text = para.text.strip()
            print(f"Paragraph text: '{text}'")  # DEBUG
            if not text:
                continue

            alignment = para.paragraph_format.alignment
            is_centered = alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
            all_bold = all(is_bold(run) for run in para.runs if run.text.strip())

            if is_centered and all_bold:
                print("Detected Topic")
                output.append({"type": "topic", "text": text})
                continue

            runs = para.runs
            first_non_empty_run_idx = next((i for i, r in enumerate(runs) if r.text.strip()), None)
            if first_non_empty_run_idx is None:
                output.append({"type": "paragraph", "text": text})
                continue

            i = first_non_empty_run_idx
            heading_runs = []
            subheading_runs = []

            while i < len(runs) and runs[i].text.strip():
                run = runs[i]
                if is_bold_and_underline(run):
                    heading_runs.append(run)
                elif is_underline(run) and not is_bold(run):
                    subheading_runs.append(run)
                else:
                    break
                i += 1

            if heading_runs:
                print("Detected Heading")
                output.append({"type": "heading", "text": text})
                continue

            if subheading_runs:
                subheading_text = "".join(run.text for run in subheading_runs).strip()
                rest_text = "".join(run.text for run in runs[i:]).strip()
                print(f"Detected Subheading: '{subheading_text}' with content: '{rest_text}'")
                entry = {"type": "subheading", "text": subheading_text}
                if rest_text:
                    entry["content"] = [{"type": "paragraph", "text": rest_text}]
                output.append(entry)
                continue

            style_name = para.style.name if para.style else ""
            if "List" in style_name or "Bullet" in style_name:
                print("Detected Bullet")
                output.append({"type": "bullet", "text": text})
            else:
                print("Detected Paragraph")
                output.append({"type": "paragraph", "text": text})

        elif block_type == 'table':
            if table_idx >= len(tables):
                print("Table index out of range!")
                break
            table = tables[table_idx]
            table_idx += 1
            table_data = []
            for row in table.rows:
                table_data.append([cell.text.strip() for cell in row.cells])
            print("Detected Table")
            output.append({"type": "table", "data": table_data})

    return output


if __name__ == "__main__":
    docx_path = "your_file.docx"
    result = extract_docx_structure(docx_path)

    import json
    with open("docx_output.json", "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)

    import pprint
    pprint.pprint(result)
